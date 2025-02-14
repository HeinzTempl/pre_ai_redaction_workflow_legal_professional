from docx import Document
import re
from llm_api import redact_text_api
from docx import Document
import re
import spacy

#def redact_text_local(text, patterns):
    #"""Ersetzt alle Vorkommen der angegebenen Regex-Muster im Text durch '[REDACTED]'."""
    #redacted_text = text
    #for pattern in patterns:
        #redacted_text = re.sub(pattern, '[REDACTED]', redacted_text)
    #return redacted_text



# Lade das spaCy-Modell (achte darauf, dass du das entsprechende deutsche Modell installiert hast)
nlp = spacy.load("de_core_news_sm")

def redact_regex(text):
    """
    Wendet Regex-Muster an, um E-Mail-Adressen und Telefonnummern zu ersetzen.
    """
    # Muster für E-Mail-Adressen
    text = re.sub(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b', '[REDACTED]', text)
    # Muster für Telefonnummern (vereinfacht)
    text = re.sub(r'\b\+?\d[\d\s-]{7,}\d\b', '[REDACTED]', text)
    return text

def redact_spacy(text):
    """
    Verwendet spaCy, um PERSON- und ORG-Entitäten zu erkennen und diese zu ersetzen.
    Personen werden durch "Person A", "Person B", ... ersetzt.
    Firmen werden durch "Firma A", "Firma B", ... ersetzt.
    """
    doc = nlp(text)
    redacted = text
    person_mapping = {}
    firm_mapping = {}
    person_counter = 0
    firm_counter = 0
    for ent in doc.ents:
        if ent.label_ == "PER":
            if ent.text not in person_mapping:
                placeholder = f"Person {chr(65 + person_counter)}"  # 65 entspricht 'A'
                person_mapping[ent.text] = placeholder
                person_counter += 1
            redacted = redacted.replace(ent.text, person_mapping[ent.text])
        elif ent.label_ == "ORG":
            if ent.text not in firm_mapping:
                placeholder = f"Firma {chr(65 + firm_counter)}"
                firm_mapping[ent.text] = placeholder
                firm_counter += 1
            redacted = redacted.replace(ent.text, firm_mapping[ent.text])
    return redacted

def process_docx(file_path, output_path):
    """
    Liest ein DOCX-Dokument, führt zuerst Regex-basierte Redaction und dann spaCy-basierte
    Redaction durch und speichert das Ergebnis in einer neuen Datei.
    """
    doc = Document(file_path)
    for para in doc.paragraphs:
        # Zuerst Regex-Redaction
        text = redact_regex(para.text)
        # Dann spaCy-Redaction für Namen/Entitäten
        text = redact_spacy(text)
        para.text = text
    doc.save(output_path)
    print(f"Redacted DOCX gespeichert: {output_path}")

# Beispielhafter Aufruf:
if __name__ == '__main__':
    input_file = 'input.docx'
    output_file = 'output_redacted.docx'
    process_docx(input_file, output_file)




def process_docx_api(file_path, output_path):
    """
    Liest ein DOCX-Dokument, extrahiert den gesamten Text und schickt ihn an die API.
    Der von der API zurückgelieferte redigierte Text wird dann in einem neuen DOCX-Dokument gespeichert.
    """
    doc = Document(file_path)
    full_text = "\n".join([para.text for para in doc.paragraphs])

    redacted_full_text = redact_text_api(full_text)

    new_doc = Document()
    # Füge jede Zeile (getrennt durch Zeilenumbrüche) als neuen Absatz ein
    for line in redacted_full_text.split("\n"):
        new_doc.add_paragraph(line)
    new_doc.save(output_path)
    print(f"API-basierte Redaktion abgeschlossen: {output_path}")

# Testweise kannst du beide Funktionen einzeln aufrufen.